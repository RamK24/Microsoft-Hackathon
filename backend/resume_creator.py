from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_heading_with_hr(doc, heading_text):
    """
    Adds a heading paragraph with a bottom border (horizontal rule)
    and minimal spacing after the heading.
    """
    p_heading = doc.add_paragraph()
    run = p_heading.add_run(heading_text)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    p_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    p_heading.paragraph_format.space_after = Pt(2)

    # Add a bottom border to the heading (acts as a horizontal rule)
    pPr = p_heading._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_header(doc, user_details):
    """
    Adds a header with:
      - A big name (all caps, navy blue, size 16)
      - On the same line below, state, country and phone separated by a vertical bar.
    """
    header_para = doc.add_paragraph()
    # Candidate Name in all caps
    name_run = header_para.add_run(user_details.get('name', '').upper() + "\n")
    name_run.font.name = 'Calibri'
    name_run.font.size = Pt(16)
    name_run.font.color.rgb = RGBColor(0, 0, 128)  # navy blue
    name_run.bold = True

    # Contact info on the same line: state, country and phone
    details = f"{user_details.get('state', '')}, {user_details.get('country', '')} | {user_details.get('phone', '')}"
    details_run = header_para.add_run(details)
    details_run.font.name = 'Calibri'
    details_run.font.size = Pt(10)
    header_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a horizontal rule after header (using a bottom border on an empty paragraph)
    hr_para = doc.add_paragraph()
    hr_para.paragraph_format.space_before = Pt(2)
    hr_para.paragraph_format.space_after = Pt(2)
    pPr = hr_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)


def create_resume(user_details, output_file='resume.docx'):
    # Create a new Document
    doc = Document()

    # Set default font style for the document (Calibri, 10)
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    # Add header with name and contact info
    add_header(doc, user_details)

    # Helper to add a section with heading and minimal spacing after content
    def add_section(heading, content_func):
        add_heading_with_hr(doc, heading.upper())
        content_func()
        # Minimal spacing after section content
        para = doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)

    # Section: Summary
    def add_summary():
        p = doc.add_paragraph(user_details.get('summary', ''))
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_after = Pt(2)

    # Section: Education (with duration on the right)
    def add_education():
        for edu in user_details.get('education', []):
            # Create a table with two columns:
            # Left: Education info, Right: Duration (right aligned)
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            edu_text = f"{edu.get('degree', '')} in {edu.get('field', '')} - {edu.get('institution', '')}"
            left_cell.text = edu_text
            right_cell.text = f"{edu.get('start_date', '')} - {edu.get('end_date', '')}"
            for para in right_cell.paragraphs:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            # Minimal spacing after each education entry
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Section: Skills
    def add_skills():
        p_skills = doc.add_paragraph(', '.join(user_details.get('skills', [])))
        p_skills.paragraph_format.line_spacing = 1.0
        p_skills.paragraph_format.space_after = Pt(2)

    # Section: Work Experience
    def add_work_experience():
        for exp in user_details.get('work_experience', []):
            # First line: Company name in bold
            p_company = doc.add_paragraph()
            run_company = p_company.add_run(exp.get('company', ''))
            run_company.bold = True
            run_company.font.name = 'Calibri'
            run_company.font.size = Pt(10)
            p_company.paragraph_format.space_after = Pt(0)

            # Second line: Use a table with two columns for job title and duration
            table = doc.add_table(rows=1, cols=2)
            table.autofit = True
            left_cell = table.rows[0].cells[0]
            right_cell = table.rows[0].cells[1]

            # Left cell: Job title (in italic)
            left_cell.text = exp.get('role', '')
            # Right cell: Duration (right aligned)
            right_cell.text = f"{exp.get('start_date', '')} - {exp.get('end_date', '')}"
            for para in right_cell.paragraphs:
                para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

            # Add summary bullet points below
            summary = exp.get('summary', '')
            if isinstance(summary, list):
                for point in summary:
                    bullet = doc.add_paragraph(point, style='List Bullet')
                    bullet.paragraph_format.line_spacing = 1.0
                    bullet.paragraph_format.space_after = Pt(0)
            else:
                bullet = doc.add_paragraph(summary, style='List Bullet')
                bullet.paragraph_format.line_spacing = 1.0
                bullet.paragraph_format.space_after = Pt(0)

            # Minimal spacing after each work experience entry
            doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Add each section
    add_section('Summary', add_summary)
    add_section('Education', add_education)
    add_section('Skills', add_skills)
    add_section('Work Experience', add_work_experience)

    # Save the document
    doc.save(output_file)
    print(f"Resume saved as {output_file}")


# if __name__ == '__main__':
#     # Example user details
#     user_details = {
#         'name': 'John Doe',
#         'state': 'California',
#         'country': 'USA',
#         'phone': '+1 234 567 890',
#         'summary': 'Experienced software developer with a passion for building scalable applications. Skilled in multiple programming languages and frameworks.',
#         'education': [
#             {
#                 'degree': 'Bachelor of Science',
#                 'field': 'Computer Science',
#                 'institution': 'University A',
#                 'start_date': '2010',
#                 'end_date': '2014'
#             }
#         ],
#         'skills': ['Python', 'Java', 'SQL', 'Git', 'REST APIs'],
#         'work_experience': [
#             {
#                 'role': 'Software Engineer',
#                 'company': 'Google',
#                 'start_date': '2018',
#                 'end_date': 'Present',
#                 'summary': [
#                     'Developed scalable backend systems.',
#                     'Optimized algorithms for performance.'
#                 ]
#             },
#             {
#                 'role': 'Senior Software Engineer',
#                 'company': 'Innovate LLC',
#                 'start_date': '2014',
#                 'end_date': '2018',
#                 'summary': ['Led a team of developers, designed scalable architecture, and improved overall system performance.']
#             }
#         ]
#     }
#
#     create_resume(user_details)
